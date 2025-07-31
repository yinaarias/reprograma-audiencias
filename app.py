import streamlit as st
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from io import BytesIO

def extract_data_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()

    data = {
        "CUI": "",
        "NI": "",
        "Acusado": "",
        "Delito": "",
        "Tipo de Procedimiento": "",
        "Fecha Original": "",
        "Hora Original": "",
        "Nueva Fecha": "",
        "Nueva Hora": "",
        "Juez": ""
    }

    # Extract data using simple string search
    lines = text.splitlines()
    for line in lines:
        if "C.U.I." in line:
            data["CUI"] = line.split(":")[-1].strip()
        elif "N.I." in line:
            data["NI"] = line.split(":")[-1].strip()
        elif "Acusado" in line:
            data["Acusado"] = line.split(":")[-1].strip()
        elif "Delito" in line:
            data["Delito"] = line.split(":")[-1].strip()
        elif "Tipo de Procedimiento" in line:
            data["Tipo de Procedimiento"] = line.split(":")[-1].strip()
        elif "programada para el día" in line or "programada para el día" in line:
            parts = line.split("programada para el")[1].split(",")
            if len(parts) >= 2:
                data["Fecha Original"] = parts[0].strip()
                data["Hora Original"] = parts[1].strip()
        elif "se fija fecha y hora para AUDIENCIA" in line:
            parts = line.split("el día")[-1].split(",")
            if len(parts) >= 2:
                data["Nueva Fecha"] = parts[0].strip()
                data["Nueva Hora"] = parts[1].strip()
        elif "Juez" in line:
            data["Juez"] = line.strip().replace("Juez", "").strip()

    return data

def generar_word(data):
    doc = Document()
    doc.add_heading("Auto de Reprogramación de Audiencia", level=1)
    doc.add_paragraph(f"C.U.I.: {data['CUI']}")
    doc.add_paragraph(f"N.I.: {data['NI']}")
    doc.add_paragraph(f"Tipo de Procedimiento: {data['Tipo de Procedimiento']}")
    doc.add_paragraph(f"Acusado: {data['Acusado']}")
    doc.add_paragraph(f"Delito: {data['Delito']}")
    doc.add_paragraph(f"Audiencia original: {data['Fecha Original']} a las {data['Hora Original']}")
    doc.add_paragraph(f"Nueva fecha de audiencia: {data['Nueva Fecha']} a las {data['Nueva Hora']}")
    doc.add_paragraph(f"Juez: {data['Juez']}")
    doc.add_paragraph("Cumplase.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_excel(data):
    df = pd.DataFrame([data])
    buffer = BytesIO()
    df.to_excel(buffer, index=False, engine='openpyxl')
    buffer.seek(0)
    return buffer

st.title("Generador de Autos de Reprogramación de Audiencia")

uploaded_file = st.file_uploader("Sube un archivo PDF de constancia judicial", type="pdf")

if uploaded_file:
    datos = extract_data_from_pdf(uploaded_file)
    st.subheader("Datos extraídos:")
    st.json(datos)

    word_file = generar_word(datos)
    excel_file = generar_excel(datos)

    st.download_button("📄 Descargar Word", word_file, file_name="Auto_Reprograma.docx")
    st.download_button("📊 Descargar Excel", excel_file, file_name="Auto_Reprograma.xlsx")